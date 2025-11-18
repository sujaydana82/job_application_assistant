import PyPDF2
from docx import Document
import pandas as pd
from typing import List, Dict
import re

class FileProcessor:
    @staticmethod
    def extract_text_from_pdf(file) -> str:
        """Extract complete text from PDF file with better parsing"""
        try:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    # Clean up the text
                    page_text = re.sub(r'\s+', ' ', page_text)  # Replace multiple spaces
                    page_text = page_text.strip()
                    text += f"--- Page {page_num + 1} ---\n{page_text}\n\n"
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file) -> str:
        """Extract complete text from DOCX file including tables"""
        try:
            doc = Document(file)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error reading DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file) -> str:
        """Extract text from TXT file"""
        try:
            content = file.read().decode('utf-8')
            return content.strip()
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
    
    @staticmethod
    def process_uploaded_file(file) -> str:
        """Process uploaded file based on its type"""
        file_type = file.type if hasattr(file, 'type') else None
        
        if file_type == "application/pdf":
            return FileProcessor.extract_text_from_pdf(file)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return FileProcessor.extract_text_from_docx(file)
        elif file_type == "text/plain":
            return FileProcessor.extract_text_from_txt(file)
        else:
            raise Exception(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def parse_cv_sections(cv_text: str) -> Dict[str, str]:
        """Parse CV into structured sections"""
        sections = {
            'personal_info': '',
            'summary': '',
            'experience': '',
            'education': '',
            'skills': '',
            'projects': '',
            'certifications': ''
        }
        
        lines = cv_text.split('\n')
        current_section = 'personal_info'
        section_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Detect section headers
            lower_line = line.lower()
            if any(keyword in lower_line for keyword in ['experience', 'work history', 'employment']):
                if section_content:
                    sections[current_section] = '\n'.join(section_content)
                current_section = 'experience'
                section_content = []
            elif any(keyword in lower_line for keyword in ['education', 'academic']):
                if section_content:
                    sections[current_section] = '\n'.join(section_content)
                current_section = 'education'
                section_content = []
            elif any(keyword in lower_line for keyword in ['skills', 'technical', 'technologies']):
                if section_content:
                    sections[current_section] = '\n'.join(section_content)
                current_section = 'skills'
                section_content = []
            elif any(keyword in lower_line for keyword in ['projects', 'portfolio']):
                if section_content:
                    sections[current_section] = '\n'.join(section_content)
                current_section = 'projects'
                section_content = []
            elif any(keyword in lower_line for keyword in ['certifications', 'certificate']):
                if section_content:
                    sections[current_section] = '\n'.join(section_content)
                current_section = 'certifications'
                section_content = []
            elif any(keyword in lower_line for keyword in ['summary', 'objective', 'about']):
                if section_content:
                    sections[current_section] = '\n'.join(section_content)
                current_section = 'summary'
                section_content = []
            else:
                section_content.append(line)
        
        # Add the last section
        if section_content:
            sections[current_section] = '\n'.join(section_content)
            
        return sections